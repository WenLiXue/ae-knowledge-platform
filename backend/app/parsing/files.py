"""文本文件提取：供本地上传和飞书 Drive 附件复用。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import shutil
import subprocess
import tempfile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_FILE_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def extract_file_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.casefold()
    if suffix == ".pdf":
        return extract_pdf_text(data)
    if suffix == ".docx":
        document = Document(BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts)
    if suffix == ".xlsx":
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    parts.append("\t".join(values))
        workbook.close()
        return "\n".join(parts)
    raise ValueError(f"unsupported file type: {suffix}")


def extract_pdf_text(data: bytes) -> str:
    """提取 PDF 文本并保留页边界，便于 RAG 引用页码。

    优先使用 pypdf 的 layout 模式，避免多栏/表格文字被完全压扁；旧版
    pypdf 不支持该参数时自动回退到普通模式。空白页不会制造无意义的
    chunk，但会保留真实页码标记，回答时可定位原文页。
    """
    reader = PdfReader(BytesIO(data))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except TypeError:  # pragma: no cover - compatibility with older pypdf
            text = page.extract_text() or ""
        text = text.replace("\x00", "").strip()
        if not text:
            continue
        pages.append(f"## 第 {page_number} 页\n\n{text}")
    if pages:
        return "\n\n".join(pages)

    # 扫描版 PDF 没有文本层时，使用容器内可选的 Tesseract OCR 降级。
    # 运行时未安装 OCR 时保持空结果，不让普通 PDF 入库流程失败。
    if not (shutil.which("pdftoppm") and shutil.which("tesseract")):
        return ""
    try:
        with tempfile.TemporaryDirectory(prefix="ae-pdf-ocr-") as workdir:
            source = Path(workdir) / "source.pdf"
            source.write_bytes(data)
            rendered = Path(workdir) / "page"
            subprocess.run(
                ["pdftoppm", "-r", "150", "-png", str(source), str(rendered)],
                check=True,
                capture_output=True,
                timeout=120,
            )
            ocr_pages: list[str] = []
            for image in sorted(Path(workdir).glob("page-*.png")):
                page_number = int(image.stem.rsplit("-", 1)[-1])
                result = subprocess.run(
                    ["tesseract", str(image), "stdout", "-l", "chi_sim+eng", "--psm", "3"],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                text = result.stdout.replace("\x00", "").strip()
                if text:
                    ocr_pages.append(f"## 第 {page_number} 页（OCR）\n\n{text}")
            return "\n\n".join(ocr_pages)
    except (OSError, subprocess.SubprocessError):
        return ""
