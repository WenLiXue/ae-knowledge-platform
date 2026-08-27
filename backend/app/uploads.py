"""本地知识文件上传：校验、文本抽取并提交现有处理流水线。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy.orm import Session

from .auth.deps import get_current_user
from .core.config import get_settings
from .db.models.user import User
from .db.session import get_db
from .knowledge import service
from .storage.local import LocalObjectStore

router = APIRouter(prefix="/api/v1/uploads", tags=["uploads"])

MAX_FILE_SIZE = 50 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def _extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.casefold()
    try:
        if suffix == ".pdf":
            return "\n\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(data)).pages)
        if suffix == ".docx":
            document = Document(BytesIO(data))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(parts)
        if suffix == ".xlsx":
            workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet in workbook.worksheets:
                parts.append(f"# {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(value.strip() for value in values):
                        parts.append("\t".join(values))
            workbook.close()
            return "\n".join(parts)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"code": "FILE_PARSE_FAILED", "message": f"无法解析文件“{filename}”。"},
        ) from exc
    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail={"code": "UNSUPPORTED_FILE_TYPE", "message": "仅支持 PDF、DOCX 和 XLSX 文件。"},
    )


@router.post("/documents", status_code=status.HTTP_202_ACCEPTED)
async def upload_documents(
    files: list[UploadFile] = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict[str, object]:
    if not files or len(files) > 20:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"code": "INVALID_FILE_COUNT", "message": "每次可上传 1–20 个文件。"},
        )

    store = LocalObjectStore(get_settings().storage_root)
    prepared: list[tuple[str, str | None, bytes, str]] = []
    for upload in files:
        filename = Path(upload.filename or "").name
        suffix = Path(filename).suffix.casefold()
        if not filename or suffix not in SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail={"code": "UNSUPPORTED_FILE_TYPE", "message": "仅支持 PDF、DOCX 和 XLSX 文件。"},
            )
        data = await upload.read(MAX_FILE_SIZE + 1)
        if not data:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail={"code": "EMPTY_FILE", "message": f"文件“{filename}”为空。"},
            )
        if len(data) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail={"code": "FILE_TOO_LARGE", "message": f"文件“{filename}”超过 50 MB。"},
            )
        text = _extract_text(filename, data).strip()
        if not text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail={"code": "EMPTY_DOCUMENT", "message": f"文件“{filename}”未提取到可用文字。"},
            )
        prepared.append((filename, upload.content_type, data, text))

    results: list[service.SubmitOutcome] = []
    for filename, content_type, data, text in prepared:
        results.append(
            service.submit_manual_upload(
                db,
                owner_user_id=user.id,
                filename=filename,
                content_type=content_type,
                data=data,
                extracted_text=text,
                store=store,
            )
        )

    return {"data": {"items": [outcome.__dict__ for outcome in results]}}
