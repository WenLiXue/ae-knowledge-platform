from __future__ import annotations

from io import BytesIO
from pathlib import Path
import shutil
import tempfile
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.main import app
from app.db.session import SessionLocal
from app.storage.local import LocalObjectStore
from app.worker.runner import WorkerRunner

client = TestClient(app)


@pytest.fixture(autouse=True)
def authenticated_client() -> None:
    client.cookies.clear()
    start = client.post("/api/v1/auth/feishu/start").json()["data"]
    response = client.get(
        f"/api/v1/auth/feishu/callback?code=auth-code&state={start['state']}",
        follow_redirects=False,
    )
    assert response.status_code == 302
    yield
    client.cookies.clear()


def _docx_bytes(text_value: str) -> bytes:
    document = Document()
    document.add_heading("测试文档", level=1)
    document.add_paragraph(text_value)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_upload_docx_creates_parse_task(monkeypatch: pytest.MonkeyPatch) -> None:
    storage_parent = Path("storage")
    storage_parent.mkdir(exist_ok=True)
    temp_root = Path(tempfile.mkdtemp(prefix="test-uploads-", dir=storage_parent))
    try:
        monkeypatch.setattr("app.uploads.get_settings", lambda: SimpleNamespace(storage_root=str(temp_root)))
        data = _docx_bytes("这是可检索的产品知识内容。")
        response = client.post(
            "/api/v1/uploads/documents",
            files=[("files", ("产品说明.docx", data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        assert response.status_code == 202, response.text
        item = response.json()["data"]["items"][0]
        assert item["duplicate"] is False
        with SessionLocal() as session:
            source_type = session.execute(
                text("SELECT source_type FROM knowledge.knowledge_sources WHERE id=:id"),
                {"id": item["source_id"]},
            ).scalar_one()
            task_type = session.execute(
                text("SELECT task_type FROM tasking.processing_tasks WHERE id=:id"),
                {"id": item["task_id"]},
            ).scalar_one()
        assert source_type == "MANUAL_UPLOAD"
        assert task_type == "PARSE"

        runner = WorkerRunner(
            worker_id="upload-test-worker",
            retry_base_delay_seconds=0,
            store=LocalObjectStore(temp_root),
        )
        for _ in range(6):
            runner.claim_and_execute(batch_size=10)
        with SessionLocal() as session:
            final_status = session.execute(
                text("SELECT status FROM knowledge.knowledge_sources WHERE id=:id"),
                {"id": item["source_id"]},
            ).scalar_one()
        assert final_status == "QUERYABLE"
    finally:
        shutil.rmtree(temp_root)


def test_upload_rejects_unsupported_file() -> None:
    response = client.post(
        "/api/v1/uploads/documents",
        files=[("files", ("payload.exe", b"not allowed", "application/octet-stream"))],
    )
    assert response.status_code == 415
    assert response.json()["detail"]["code"] == "UNSUPPORTED_FILE_TYPE"
